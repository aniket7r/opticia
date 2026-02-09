"""Report generation service.

Generates structured documents (reports, summaries, comparisons, etc.)
using a separate Gemini text API call (non-live), enriched with
web search results and conversation context.
"""

import io
import logging
import uuid
from typing import Any

from google import genai
from google.genai import types

from app.core.config import settings

logger = logging.getLogger(__name__)

# Gemini text model for report generation (non-live, non-audio)
REPORT_MODEL = "gemini-2.5-flash"


async def _search_topic(topic: str, num_queries: int = 3) -> str:
    """Run web searches using Google Search grounding and return formatted results."""
    from app.services.tools.web_search import web_search_handler

    queries = [topic]
    if num_queries >= 2:
        queries.append(f"{topic} detailed information")
    if num_queries >= 3:
        queries.append(f"{topic} comparison guide")

    all_answers: list[str] = []
    all_sources: list[dict[str, str]] = []
    seen_urls: set[str] = set()

    for query in queries:
        try:
            result = await web_search_handler({"query": query})
            if result.success and result.result:
                answer = result.result.get("answer", "")
                if answer:
                    all_answers.append(f"[{query}]: {answer}")
                for s in result.result.get("sources", []):
                    url = s.get("url", "")
                    if url and url not in seen_urls:
                        seen_urls.add(url)
                        all_sources.append(s)
        except Exception as e:
            logger.warning(f"Search failed for '{query}': {e}")

    if not all_answers:
        return "(No web search results available)"

    formatted = "\n\n".join(all_answers)
    if all_sources:
        source_lines = "\n".join(f"- {s.get('title', '')}: {s.get('url', '')}" for s in all_sources[:8])
        formatted += f"\n\nSources:\n{source_lines}"

    return formatted


def _build_generation_prompt(
    topic: str,
    context_history: list[dict[str, str]],
    search_results: str,
) -> str:
    """Build the prompt for the Gemini text API report generation call."""
    # Build conversation context (last relevant exchanges, max 3000 chars)
    context_lines = []
    total_chars = 0
    for entry in reversed(context_history):
        role = entry.get("role", "unknown")
        content = entry.get("content", "")[:500]
        line = f"[{role}]: {content}"
        if total_chars + len(line) > 3000:
            break
        context_lines.insert(0, line)
        total_chars += len(line)

    context_text = "\n".join(context_lines) if context_lines else "(No conversation context)"

    return f"""You are a professional document generator. Create a well-structured document based on the topic, conversation context, and web research provided below.

TOPIC: {topic}

CONVERSATION CONTEXT (what the user discussed):
{context_text}

WEB RESEARCH (additional information gathered):
{search_results}

INSTRUCTIONS:
1. Create a comprehensive, well-structured document about the topic.
2. Choose the most appropriate format based on the content:
   - **Report**: For informational topics with sections and findings
   - **Comparison Table**: For comparing items (medicines, products, options)
   - **Summary Notes**: For condensing a discussion into key takeaways
   - **Checklist**: For actionable step-by-step guides
   - **Research Brief**: For technical or research topics
3. Use clear Markdown formatting: headings (##, ###), bullet points, tables, bold text.
4. Include information from BOTH the conversation context AND web research.
5. Add source references where applicable.
6. Keep the document focused, practical, and useful as a reference.
7. Length should match the depth of the topic — be thorough but not paddy.

Generate the document now in Markdown format:"""


async def generate_report(
    topic: str,
    context_history: list[dict[str, str]],
    session_id: str,
) -> dict[str, Any]:
    """Orchestrate the full report generation pipeline.

    Returns dict with keys: reportId, topic, markdownContent, htmlContent
    """
    report_id = str(uuid.uuid4())
    logger.info(f"Starting report generation: '{topic}' (report_id={report_id})")

    try:
        # Step 1: Web searches
        search_results = await _search_topic(topic)
        logger.info(f"Search complete for report {report_id}")

        # Step 2: Build prompt
        prompt = _build_generation_prompt(topic, context_history, search_results)

        # Step 3: Call Gemini text API (non-live, async)
        client = genai.Client(api_key=settings.gemini_api_key)
        response = await client.aio.models.generate_content(
            model=REPORT_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.7,
                max_output_tokens=4096,
            ),
        )

        markdown_content = response.text or ""
        if not markdown_content.strip():
            raise ValueError("Gemini returned empty response")

        logger.info(f"Report generated: {report_id} ({len(markdown_content)} chars)")

        # Step 4: Convert markdown to HTML for preview
        import markdown
        html_content = markdown.markdown(
            markdown_content,
            extensions=["tables", "fenced_code", "nl2br"],
        )

        return {
            "reportId": report_id,
            "topic": topic,
            "markdownContent": markdown_content,
            "htmlContent": html_content,
        }

    except Exception as e:
        logger.error(f"Report generation failed for {report_id}: {e}", exc_info=True)
        raise


def _clean_text(text: str) -> str:
    """Clean text for XML compatibility — remove control chars, strip markdown bold."""
    import re
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    # Remove control characters that XML doesn't accept
    text = "".join(c for c in text if ord(c) >= 32 or c in "\n\r\t")
    return text


def export_docx(markdown_content: str, topic: str) -> io.BytesIO:
    """Convert markdown to DOCX using python-docx."""
    from docx import Document
    from docx.shared import Pt
    import re

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(11)

    lines = markdown_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        # Table detection
        elif "|" in line and i + 1 < len(lines) and "---" in lines[i + 1]:
            # Parse markdown table
            headers = [c.strip() for c in line.split("|") if c.strip()]
            i += 1  # Skip separator line
            rows = []
            while i + 1 < len(lines) and "|" in lines[i + 1]:
                i += 1
                cells = [c.strip() for c in lines[i].split("|") if c.strip()]
                rows.append(cells)

            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Table Grid"
                # Headers
                for j, h in enumerate(headers):
                    table.rows[0].cells[j].text = _clean_text(h)
                # Data rows
                for r_idx, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        if j < len(headers):
                            table.rows[r_idx + 1].cells[j].text = _clean_text(cell)
        # Bullet points
        elif line.startswith("- ") or line.startswith("* "):
            text = _clean_text(line[2:].strip())
            doc.add_paragraph(text, style="List Bullet")
        # Numbered list
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line).strip()
            text = _clean_text(text)
            doc.add_paragraph(text, style="List Number")
        # Empty line
        elif not line.strip():
            pass
        # Regular paragraph
        else:
            text = _clean_text(line.strip())
            if text:
                doc.add_paragraph(text)

        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_pdf(html_content: str, topic: str) -> io.BytesIO:
    """Convert HTML to PDF using weasyprint."""
    from weasyprint import HTML

    # Wrap in a styled HTML document
    styled_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
           font-size: 11pt; line-height: 1.6; color: #333; max-width: 800px; margin: 0 auto; padding: 40px; }}
    h1 {{ color: #1a1a1a; border-bottom: 2px solid #e5e5e5; padding-bottom: 8px; }}
    h2 {{ color: #2a2a2a; margin-top: 24px; }}
    h3 {{ color: #3a3a3a; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
    th {{ background-color: #f5f5f5; font-weight: 600; }}
    tr:nth-child(even) {{ background-color: #fafafa; }}
    ul, ol {{ padding-left: 24px; }}
    li {{ margin-bottom: 4px; }}
    code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; }}
    pre {{ background: #f4f4f4; padding: 16px; border-radius: 6px; overflow-x: auto; }}
</style>
</head>
<body>
{html_content}
</body>
</html>"""

    buffer = io.BytesIO()
    HTML(string=styled_html).write_pdf(buffer)
    buffer.seek(0)
    return buffer
